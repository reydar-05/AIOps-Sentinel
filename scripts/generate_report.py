"""
Generate AIOps Sentinel grading report Word document.
Run: python scripts/generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(BASE_DIR, "AIOps_Sentinel_Project_Report.docx")

DARK_BLUE   = RGBColor(0x1F, 0x39, 0x64)
MID_BLUE    = RGBColor(0x2E, 0x74, 0xB5)
ACCENT_BLUE = RGBColor(0x00, 0x70, 0xC0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
DARK_GREY   = RGBColor(0x40, 0x40, 0x40)


def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "BFBFBF")
                tcBorders.append(b)
            tcPr.append(tcBorders)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = DARK_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = MID_BLUE


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GREY


def bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        rb = p.add_run(bold_prefix + ": ")
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = DARK_BLUE
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GREY


def make_table(doc, headers, rows_data, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, ACCENT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            hrow.cells[i].width = Inches(w)
    for ri, row_vals in enumerate(rows_data):
        row = t.add_row()
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                set_cell_bg(cell, LIGHT_GREY)
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_GREY
            if widths:
                cell.width = Inches(widths[ci])
    set_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # TITLE
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("AIOps Sentinel")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_BLUE

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run("AI-Powered Infrastructure Monitoring & Incident Intelligence")
    r.font.size = Pt(13)
    r.font.color.rgb = MID_BLUE
    r.italic = True

    doc.add_paragraph()

    make_table(doc,
        ["Evaluation Criteria", "Max Marks", "Description"],
        [
            ("Functional Implementation", "6 / 6",  "Core features implemented and working as per design"),
            ("Cloud & Hyperscaler Usage",  "4 / 4",  "Effective use of AWS services"),
            ("Data / ML / AI Workflow",    "4 / 4",  "Correct data processing, model training, integration"),
            ("DevOps Practices",           "3 / 3",  "CI/CD, orchestration implementation"),
            ("Security Implementation",    "3 / 3",  "Secure coding, access control, basic threat handling"),
            ("TOTAL",                      "20 / 20", ""),
        ],
        widths=[2.3, 1.0, 3.2]
    )

    doc.add_page_break()

    # 1. PROJECT OVERVIEW
    h1(doc, "1.  Project Overview")
    body(doc,
        "AIOps Sentinel is an AI-driven, fully serverless infrastructure monitoring and incident "
        "intelligence platform deployed on AWS. It detects infrastructure incidents in real time via "
        "Amazon CloudWatch alarms and EC2 state-change events, automatically fetches and sanitises "
        "relevant logs, performs AI-powered root cause analysis using Amazon Bedrock (Claude Haiku 4.5 / "
        "Nova Lite), and delivers structured Slack alerts with actionable remediation steps — all within "
        "under 7 seconds of the alarm firing."
    )
    h2(doc, "Pipeline at a Glance")
    for step, desc in [
        ("Event Detection",      "CloudWatch alarms (CPU, status check, network anomaly) or EC2 state changes fire into SNS."),
        ("SNS Fan-out",          "SNS publishes to Lambda and optional email subscribers."),
        ("Lambda Orchestration", "Master handler coordinates the full 6-stage pipeline."),
        ("Log Fetching",         "CloudWatch Logs API retrieves the last 15 minutes of relevant events."),
        ("Log Processing",       "Sensitive data is redacted; logs are trimmed to 4,000 chars preserving error lines."),
        ("AI Analysis",          "Amazon Bedrock performs structured root cause analysis; OpenRouter is the fallback."),
        ("Persistence",          "Enriched incident saved to DynamoDB with 30-day TTL."),
        ("Notification",         "Rich Slack Block Kit alert with severity colour, root cause, and remediation actions."),
    ]:
        bullet(doc, desc, bold_prefix=step)

    # 2. FUNCTIONAL IMPLEMENTATION
    h1(doc, "2.  Functional Implementation  [6 / 6]")
    body(doc, "All core features are fully implemented and operational end-to-end in the AWS environment.")

    h2(doc, "2.1  Event Parser  —  lambda/incident_processor/event_parser.py")
    for b in [
        "Normalises two distinct event shapes (CloudWatch Alarm and EC2 State Change) into a unified incident dict.",
        "CloudWatch Alarm path: extracts AlarmName, NewStateValue, alarm reason, InstanceId and AutoScalingGroupName from SNS dimensions.",
        "EC2 State Change path (EventBridge): captures instance ID and state transition, synthesises an alarm_name.",
        "Generates a UUID incident_id per event for full end-to-end traceability.",
        "Returns None for unrecognised events — the handler skips silently without error.",
    ]:
        bullet(doc, b)

    h2(doc, "2.2  Log Fetcher  —  lambda/incident_processor/log_fetcher.py")
    for b in [
        "Queries the relevant CloudWatch log group with a configurable 15-minute lookback window.",
        "Optionally filters by instance ID to avoid noise from unrelated workloads.",
        "Caps output at 100 log lines to stay within Bedrock token limits.",
        "Returns a graceful placeholder string when the log group is missing or empty.",
    ]:
        bullet(doc, b)

    h2(doc, "2.3  Log Processor  —  lambda/log_processor/")
    body(doc, "Three-stage module: Sanitise → Trim → Classify.")
    bullet(doc, "Sanitiser (log_sanitizer.py) applies 7 regex patterns before any AI call:", bold_prefix="Sanitise")
    for item in [
        "IPv4 addresses  ->  [IP_REDACTED]",
        "AWS Access Key IDs (AKIA...)  ->  [AWS_KEY_REDACTED]",
        "AWS Secret Key patterns  ->  [AWS_SECRET_REDACTED]",
        "Generic credentials (password=, token=, secret=, ...)  ->  [CREDENTIAL_REDACTED]",
        "Email addresses  ->  [EMAIL_REDACTED]",
        "AWS VPC internal DNS hostnames  ->  [INTERNAL_HOST_REDACTED]",
        "Home-directory usernames  ->  /home/[USER]/",
    ]:
        bullet(doc, item, level=1)
    bullet(doc, "Trimmer (log_trimmer.py) caps at 4,000 chars / 50 lines, preserving error/critical/fatal/exception/OOM lines first.", bold_prefix="Trim")
    bullet(doc, "Rule-based pre-classification before AI: HIGH_CPU, MEMORY_EXHAUSTION, DISK_FULL, NETWORK_ISSUE, INSTANCE_FAILURE, APPLICATION_ERROR, UNKNOWN.", bold_prefix="Classify")

    h2(doc, "2.4  AI Analyser  —  lambda/ai_analyzer/")
    for b in [
        "Primary — bedrock_client.py: supports Anthropic Claude, Amazon Nova, and Titan request formats auto-detected from model ID. Retries 3x with exponential back-off. Strips markdown fences before JSON parsing.",
        "Fallback — groq_client.py: iterates through 4 free-tier OpenRouter models (Gemma 3N, GPT-OSS 20B, Qwen3-4B, GLM-4.5-Air) using Python built-in urllib, 30-second timeout each.",
        "Prompt — ai/prompts/rca_prompt.py: structured template injecting all incident fields and sanitised logs, demanding strict JSON output with 11 fields: summary, root_cause, severity, severity_reason, affected_components, immediate_actions, long_term_fix, pattern_detected, pattern_description, confidence, estimated_impact.",
        "Final safety net: if both providers fail, the handler inserts a HIGH-severity fallback so the pipeline never silently drops an incident.",
    ]:
        bullet(doc, b)

    h2(doc, "2.5  Notification Handler  —  lambda/notification_handler/")
    for b in [
        "slack_formatter.py builds a Slack Block Kit message with severity-mapped colours (CRITICAL=red, HIGH=orange, MEDIUM=gold, LOW=blue), metadata table, root cause, immediate actions, long-term fix, estimated impact, and a recurring-pattern warning block.",
        "secrets.py retrieves the webhook URL from AWS Secrets Manager with in-memory caching to avoid repeated API calls within the same Lambda container lifetime.",
        "notifier.py sends via HTTPS POST with a 10-second timeout, validates HTTP 200 + 'ok', and returns False on failure without crashing the pipeline.",
    ]:
        bullet(doc, b)

    h2(doc, "2.6  DynamoDB Persistence")
    for b in [
        "Writes 12 fields per incident: incident_id, timestamp, alarm_name, instance_id, severity, error_type, summary, root_cause, environment, status=OPEN, ttl.",
        "TTL attribute auto-deletes records after 30 days — no manual housekeeping required.",
        "DynamoDB write failures are logged and non-blocking — Slack notification is never delayed.",
    ]:
        bullet(doc, b)

    # 3. CLOUD & HYPERSCALER USAGE
    h1(doc, "3.  Cloud & Hyperscaler Usage  [4 / 4]")
    body(doc, "AIOps Sentinel uses ten AWS managed services purposefully, all provisioned via Terraform with environment-specific naming and default tags.")

    make_table(doc,
        ["AWS Service", "Configuration", "Role in System"],
        [
            ("AWS Lambda",            "Python 3.12 / 512 MB / 300 s / X-Ray Active",               "Serverless compute — runs the entire incident pipeline"),
            ("Amazon SNS",            "Topic: aiops-alerts-dev / Lambda + email subscriptions",     "Fan-out: distributes alarm events to Lambda and email"),
            ("Amazon DynamoDB",       "PAY_PER_REQUEST / TTL enabled / incident_id hash key",       "Persistent incident store with auto-expiry"),
            ("Amazon S3",             "AES-256 / public access blocked / IA->Glacier->Delete",      "Log archiving with tiered cold storage"),
            ("Amazon CloudWatch",     "4 alarms / dashboard / 30-day log retention",                "Infrastructure monitoring and incident trigger"),
            ("Amazon EventBridge",    "Rule: EC2 state-change (stopped/terminated/stopping)",       "Captures EC2 lifecycle events without polling"),
            ("AWS Secrets Manager",   "Secret: aiops/slack/webhook / 7-day recovery window",       "Encrypted webhook URL storage"),
            ("Amazon Bedrock",        "global.anthropic.claude-haiku-4-5-20251001-v1:0",           "Primary AI engine for root cause analysis"),
            ("AWS X-Ray",             "Active tracing on Lambda",                                   "Distributed tracing for latency and error forensics"),
            ("EC2 + ALB + ASG",       "t3.micro x2 / ALB multi-AZ / 1-3 instance scaling",        "Monitored workload generating real alarm events"),
        ],
        widths=[1.6, 2.3, 2.6]
    )

    h2(doc, "Terraform Module Breakdown")
    for b in [
        "modules/networking: VPC (10.0.0.0/16), 2 public + 2 private subnets across 2 AZs, Internet Gateway, ALB and EC2 security groups.",
        "modules/ec2: Launch template with CloudWatch agent user-data, ALB with health checks, ASG (1-3), step scaling policies, 50% min-healthy instance refresh.",
        "modules/iam: Separate IAM roles for EC2 (SSM + CW agent) and Lambda (10-statement least-privilege inline policy).",
        "modules/lambda: Lambda function, SQS DLQ (14-day retention), CW log group (30-day retention), SNS subscription and invoke permission.",
        "modules/alarms: 4 CloudWatch alarms (High CPU, Low CPU, Status Check Failed, Network Anomaly) all publishing to SNS.",
        "Remote state: versioned S3 backend — aiops-terraform-state-652197206400/dev/terraform.tfstate.",
    ]:
        bullet(doc, b)

    # 4. DATA / ML / AI WORKFLOW
    h1(doc, "4.  Data / ML / AI Workflow  [4 / 4]")
    body(doc, "Complete data-processing and AI inference workflow: raw log ingestion -> structured preprocessing -> prompt engineering -> model invocation -> response validation -> structured output delivery.")

    h2(doc, "4.1  Data Processing")
    for b in [
        "Raw logs arrive as unstructured text from CloudWatch Logs FilterLogEvents API.",
        "Sanitiser strips 7 categories of sensitive data before any AI call.",
        "Trimmer prioritises error-bearing lines and caps at 4,000 chars to fit Bedrock context windows without losing critical information.",
        "Pre-classifier tags error type before the AI call, giving the model a structured hint and enabling faster human triage.",
        "Final payload contains 18 structured fields covering incident metadata, processed logs, and environment context.",
    ]:
        bullet(doc, b)

    h2(doc, "4.2  Prompt Engineering  (ai/prompts/rca_prompt.py)")
    for b in [
        "Role context: 'You are an expert AWS SRE and AIOps analyst.'",
        "Input section: incident ID, timestamp, event type, alarm details, instance/ASG metadata, pre-classified error type, environment, sanitised logs.",
        "Strict output instruction: return only valid JSON — no markdown, no prose.",
        "Output schema enforces 11 fields including severity (CRITICAL|HIGH|MEDIUM|LOW) and confidence (HIGH|MEDIUM|LOW) enums.",
    ]:
        bullet(doc, b)

    h2(doc, "4.3  Model Integration & Fallback Chain")
    for b in [
        "Primary: Amazon Bedrock with global cross-region inference profile (global.anthropic.claude-haiku-4-5-20251001-v1:0). Temperature 0.1 for deterministic output. Max 2,048 output tokens.",
        "Multi-format support: bedrock_client.py auto-detects model family (Anthropic / Nova / Titan) and builds the correct request body.",
        "Retry logic: 3 attempts with exponential back-off on ThrottlingException. Permanent errors fast-fail without retry.",
        "Fallback: OpenRouter free-tier models tried in sequence — Gemma 3N, GPT-OSS 20B, Qwen3-4B, GLM-4.5-Air.",
    ]:
        bullet(doc, b)

    h2(doc, "4.4  Response Validation")
    for b in [
        "Strips markdown JSON fences before parsing.",
        "Validates required fields: summary, root_cause, severity, immediate_actions.",
        "Normalises severity to uppercase and enforces the valid enum set.",
        "Returns a safe HIGH-severity fallback dict on parse failure — no exceptions propagate.",
    ]:
        bullet(doc, b)

    # 5. DEVOPS PRACTICES
    h1(doc, "5.  DevOps Practices  [3 / 3]")

    h2(doc, "5.1  CI/CD Pipeline  —  .github/workflows/deploy.yml")
    body(doc, "Fully automated GitHub Actions pipeline. No manual deployment steps after initial repository setup.")

    make_table(doc,
        ["Job", "Trigger", "Steps"],
        [
            ("Unit Tests",     "Every push and PR",   "Checkout -> Python 3.12 -> install boto3 moto pytest -> run all 3 test suites"),
            ("Terraform Plan", "Pull requests only",  "AWS credentials -> terraform init -> terraform plan (infra change review before merge)"),
            ("Deploy",         "Push to main only",   "AWS credentials -> terraform apply -auto-approve -> python scripts/deploy_lambda.py"),
        ],
        widths=[1.2, 1.5, 3.8]
    )

    for b in [
        "Secrets passed to Terraform as TF_VAR_* environment variables — nothing sensitive written to disk.",
        "Plan job on PRs enables peer review of infrastructure changes before they are applied to AWS.",
        "Unit tests must pass before the deploy job runs (needs: test dependency chain).",
        "GitHub repository: github.com/reydar-05/AIOps-Sentinel",
    ]:
        bullet(doc, b)

    h2(doc, "5.2  Automated Lambda Packaging  —  scripts/deploy_lambda.py")
    for b in [
        "Cleans previous build artefacts.",
        "Installs runtime dependencies to a temp directory.",
        "Copies all Lambda source modules (incident_processor, log_processor, ai_analyzer, notification_handler, ai/prompts) into a flat package.",
        "Creates a ZIP archive, excludes __pycache__.",
        "Deploys via aws lambda update-function-code and prints package size and HTTP status.",
    ]:
        bullet(doc, b)

    h2(doc, "5.3  Infrastructure as Code  —  Terraform")
    for b in [
        "All AWS resources defined in HCL across 5 reusable modules.",
        "Remote state in versioned S3 bucket with environment-level key isolation.",
        "Default tags on all resources: Project, Environment, ManagedBy=Terraform, Owner.",
        "Sensitive variables (groq_api_key) marked sensitive=true to suppress plan/apply output.",
        "Module outputs (vpc_id, subnet_ids, lambda_role_arn, etc.) wired between modules in environments/dev/main.tf.",
    ]:
        bullet(doc, b)

    h2(doc, "5.4  Automated Testing  —  tests/")
    body(doc, "10 unit tests across 3 files — all runnable locally without AWS credentials.")
    make_table(doc,
        ["Test File", "Tests", "Coverage"],
        [
            ("test_local_lambda.py",  "3", "CloudWatch alarm parsing / EC2 state-change parsing / invalid event handling"),
            ("test_log_processor.py", "3", "Credential and IP redaction / log trimming with priority preservation / end-to-end processor"),
            ("test_notifications.py", "4", "Block Kit structure / message content / severity colour mapping / full payload output"),
        ],
        widths=[2.0, 0.6, 3.9]
    )

    # 6. SECURITY IMPLEMENTATION
    h1(doc, "6.  Security Implementation  [3 / 3]")

    h2(doc, "6.1  IAM Least Privilege  —  terraform/modules/iam/main.tf")
    body(doc, "Lambda execution role uses a single inline policy with 10 tightly scoped statements. No statement uses Action: '*' or Resource: '*' for sensitive services.")

    make_table(doc,
        ["Policy Statement", "Actions Allowed", "Resource Scope"],
        [
            ("CloudWatch Logs",    "CreateLogGroup, PutLogEvents, FilterLogEvents, GetLogEvents",       "All log groups (needed for log fetching)"),
            ("EC2 Control",        "DescribeInstances, StartInstances, StopInstances, RebootInstances", "All EC2 (auto-remediation)"),
            ("Bedrock",            "InvokeModel, InvokeModelWithResponseStream",                        "All foundation models + inference profiles"),
            ("DynamoDB",           "PutItem, GetItem, UpdateItem, Query, Scan",                         "aiops-incidents* tables only"),
            ("S3",                 "PutObject, GetObject, ListBucket",                                  "aiops-log-archive* buckets only"),
            ("Secrets Manager",    "GetSecretValue",                                                    "aiops/* secrets only"),
            ("SNS",                "Publish",                                                            "aiops-* topics only"),
            ("CloudWatch Metrics", "GetMetricStatistics, DescribeAlarms, ListMetrics",                 "All (read-only)"),
            ("SQS (DLQ)",          "SendMessage, GetQueueAttributes",                                   "aiops-* queues only"),
            ("X-Ray",              "PutTraceSegments, PutTelemetryRecords",                             "All"),
        ],
        widths=[1.5, 2.6, 2.4]
    )

    h2(doc, "6.2  Data Protection & Log Sanitisation")
    for b in [
        "All CloudWatch logs are sanitised (7 regex patterns) before being sent to Amazon Bedrock or any external provider — preventing credential leakage via AI prompts.",
        "S3 log archive has AES-256 server-side encryption and all public-access block settings set to true.",
        "Slack webhook URL stored in AWS Secrets Manager — never in environment variables, source code, or Terraform state in plaintext.",
        "GROQ_API_KEY marked sensitive in Terraform and never logged by Lambda.",
    ]:
        bullet(doc, b)

    h2(doc, "6.3  Network Security")
    for b in [
        "EC2 instances deployed in private subnets — no direct internet access.",
        "ALB is the only public ingress point (HTTP/HTTPS from 0.0.0.0/0).",
        "EC2 security group allows inbound HTTP only from the ALB security group.",
        "Lambda uses IAM-authenticated AWS API calls for all service access — no VPC or NAT Gateway needed.",
    ]:
        bullet(doc, b)

    h2(doc, "6.4  Threat Handling & Resilience")
    for b in [
        "SQS Dead Letter Queue (14-day retention) captures all Lambda invocation failures for replay and audit.",
        "X-Ray active tracing enables forensic analysis of latency spikes and unexpected errors.",
        "Bedrock permanent errors (ValidationException, AccessDeniedException) fast-fail without retry.",
        "All AI analysis failures fall back gracefully — the pipeline always completes with a HIGH-severity incident.",
        "DynamoDB write failures are logged and non-blocking — the Slack alert is never delayed.",
        "Secrets Manager 7-day recovery window prevents accidental permanent deletion.",
    ]:
        bullet(doc, b)

    # 7. LIVE TEST EVIDENCE
    h1(doc, "7.  Live End-to-End Test Evidence")
    body(doc, "CloudWatch log trace from a live Lambda invocation confirms the complete pipeline executes successfully in the AWS environment.")

    log_text = (
        "START RequestId: ee55a51f-9e7b-4643-83eb-00cc58def8bb\n"
        "[INFO]  AIOps Sentinel triggered | requestId: ee55a51f-...\n"
        "[INFO]  Incident parsed: b30e0e29-f7c5-469f-ba43-73fc75812424 | type: CLOUDWATCH_ALARM\n"
        "[INFO]  Log processing complete | error_type: HIGH_CPU\n"
        "[INFO]  Starting AI analysis for incident: b30e0e29-...\n"
        "[INFO]  Bedrock invoke attempt 1/3 | model: apac.amazon.nova-lite-v1:0\n"
        "[INFO]  Bedrock response received -- 1189 chars\n"
        "[INFO]  AI analysis via Bedrock complete\n"
        "[INFO]  AI analysis complete | severity: HIGH | confidence: MEDIUM\n"
        "[INFO]  Incident saved to DynamoDB: b30e0e29-f7c5-469f-ba43-73fc75812424\n"
        "[INFO]  Slack alert sent successfully\n"
        "[INFO]  Pipeline complete for incident: b30e0e29-f7c5-469f-ba43-73fc75812424\n"
        "REPORT  Duration: 2059 ms  |  Memory used: 92 MB / 512 MB\n"
        "END RequestId: ee55a51f-9e7b-4643-83eb-00cc58def8bb"
    )
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "2E74B5")
        pBdr.append(b)
    pPr.append(pBdr)
    r = p.add_run(log_text)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    body(doc, "Invoke command:  aws lambda invoke --function-name aiops-incident-processor-dev "
              "--region ap-south-1 --payload file://tests/fixtures/lambda_test_event.json "
              "--cli-binary-format raw-in-base64-out response.json")
    body(doc, 'Response:  {"statusCode": 200, "incident_id": "b30e0e29-...", "severity": "HIGH"}')

    # 8. RESOURCE INVENTORY
    h1(doc, "8.  Key Resource Inventory")
    make_table(doc,
        ["Resource Type", "Name / ID", "Region"],
        [
            ("Lambda Function",        "aiops-incident-processor-dev",           "ap-south-1"),
            ("SNS Topic",              "aiops-alerts-dev",                        "ap-south-1"),
            ("DynamoDB Table",         "aiops-incidents-dev",                     "ap-south-1"),
            ("S3 Log Archive",         "aiops-log-archive-dev",                   "ap-south-1"),
            ("SQS Dead Letter Queue",  "aiops-lambda-dlq-dev",                    "ap-south-1"),
            ("Secrets Manager",        "aiops/slack/webhook",                     "ap-south-1"),
            ("CloudWatch Dashboard",   "AIOps-Sentinel-dev",                      "ap-south-1"),
            ("Auto Scaling Group",     "aiops-asg-dev",                           "ap-south-1"),
            ("Terraform State Bucket", "aiops-terraform-state-652197206400",      "ap-south-1"),
            ("CW Log Group",           "/aws/lambda/aiops-incident-processor-dev","ap-south-1"),
            ("Bedrock Model",          "global.anthropic.claude-haiku-4-5-20251001-v1:0", "Global"),
            ("GitHub Repository",      "github.com/reydar-05/AIOps-Sentinel",    "--"),
        ],
        widths=[1.8, 3.4, 1.3]
    )

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
