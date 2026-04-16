"""
AIOps Sentinel — Review Presentation Document Generator
Generates AIOps_Sentinel_Review.docx at the project root.
Run: python scripts/generate_doc.py
"""

import subprocess
import sys
import os

# ── Auto-install python-docx if missing ───────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h


def add_table(doc, headers, rows, header_color="1F3864"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        tbl_row = table.rows[r_idx + 1]
        bg = "EBF0FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tbl_row.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_para(doc, text, bold=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Document Build ─────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ── COVER PAGE ────────────────────────────────────────────────
    doc.add_paragraph("\n\n")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AIOps Sentinel")
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("AI-Driven Automated Infrastructure Incident Intelligence")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("\n")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run("Environment: DEV   |   Region: ap-south-1   |   Account: 652197206400")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph("\n")
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = tag.add_run("Powered by Amazon Bedrock + OpenRouter  |  Deployed via Terraform + GitHub Actions")
    r4.font.size = Pt(10)
    r4.font.italic = True
    r4.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1, "1F3864")
    add_divider(doc)
    add_para(doc,
        "AIOps Sentinel is a fully automated, AI-powered infrastructure monitoring and incident "
        "intelligence platform deployed on AWS. It continuously monitors EC2 infrastructure via "
        "CloudWatch alarms and EventBridge events. When an anomaly is detected — such as CPU "
        "spiking above 80%, a status check failure, or an instance terminating — the platform "
        "automatically fetches and sanitizes relevant logs, runs an AI root cause analysis using "
        "Amazon Bedrock (Claude 3.5 Sonnet), and delivers a structured Slack alert with the "
        "diagnosis, severity, immediate actions, and long-term fix — all within 5 to 7 seconds "
        "of the alarm firing."
    )
    doc.add_paragraph()
    add_para(doc, "Key Outcomes:", bold=True)
    add_bullet(doc, "Alarm-to-Slack-alert in ~5-7 seconds (fully automated, zero human intervention)")
    add_bullet(doc, "AI-generated root cause analysis with severity, affected components, and action plan")
    add_bullet(doc, "All incidents persisted in DynamoDB with 30-day retention for audit and trend analysis")
    add_bullet(doc, "Dual AI fallback: Bedrock primary, OpenRouter free-tier models as backup")
    add_bullet(doc, "Fully codified infrastructure (Terraform) with 4-stage CI/CD (GitHub Actions)")

    doc.add_page_break()

    # ── 2. SYSTEM ARCHITECTURE ────────────────────────────────────
    add_heading(doc, "2. System Architecture", 1, "1F3864")
    add_divider(doc)
    add_para(doc,
        "AIOps Sentinel is a serverless, event-driven architecture built entirely on AWS managed "
        "services. There are no long-running servers in the processing pipeline — everything is "
        "triggered by events and scales automatically."
    )
    doc.add_paragraph()
    add_heading(doc, "AWS Services Used", 2, "2E74B5")
    add_table(doc,
        ["Service", "Name / Config", "Role"],
        [
            ["AWS Lambda", "aiops-incident-processor-dev\nPython 3.12 | 512 MB | 300s timeout", "Full incident pipeline: parse → log fetch → AI analysis → notify"],
            ["Amazon SNS", "aiops-alerts-dev", "Fan-out hub: receives alarms from CloudWatch + EventBridge, triggers Lambda"],
            ["Amazon EventBridge", "aiops-ec2-state-change-dev", "Captures EC2 state changes (stopped/terminated/stopping) → SNS"],
            ["Amazon Bedrock", "claude-3-5-sonnet-20241022-v2:0", "Primary AI — structured root cause analysis (RCA) via prompt engineering"],
            ["OpenRouter API", "Gemma 3N / Nemotron / Qwen / GPT-OSS\n(free tier)", "Fallback AI — invoked if Bedrock is unavailable; tries 5 models in sequence"],
            ["Amazon DynamoDB", "aiops-incidents-dev\nPay-per-request | 30-day TTL", "Incident store — persists every processed incident for audit and querying"],
            ["Amazon S3", "aiops-log-archive-dev\n30d→IA, 90d→Glacier, 365d→delete", "Log archive — stores raw logs with tiered lifecycle policy"],
            ["AWS Secrets Manager", "aiops/slack/webhook", "Securely stores Slack webhook URL; fetched at Lambda runtime"],
            ["Amazon SQS (DLQ)", "aiops-lambda-dlq-dev\n14-day retention", "Dead letter queue — captures failed Lambda invocations for replay/debug"],
            ["Amazon CloudWatch", "Alarms, Dashboard, Logs, X-Ray", "Monitoring: 4 metric alarms, live dashboard, log retention (30d), distributed tracing"],
            ["Amazon EC2 + ASG", "t3.micro | 1-3 instances (desired: 2)\naiops-asg-dev", "Monitored infrastructure — the target being watched and protected"],
            ["Application Load Balancer", "aiops-alb-dev (public, multi-AZ)", "Distributes traffic across ASG instances; health check target"],
            ["VPC + Subnets", "10.0.0.0/16\n2 public + 2 private subnets", "Network isolation; private subnets for EC2, public for ALB"],
            ["AWS IAM", "Lambda execution role\nEC2 instance profile", "Least-privilege roles scoped to aiops-* resources only"],
            ["Terraform S3 Backend", "aiops-terraform-state-*", "Shared remote state storage for Terraform with versioning"],
        ]
    )

    doc.add_page_break()

    # ── 3. END-TO-END EVENT FLOW ──────────────────────────────────
    add_heading(doc, "3. End-to-End Event Flow", 1, "1F3864")
    add_divider(doc)
    add_para(doc,
        "The entire pipeline is event-driven. A single infrastructure anomaly triggers an "
        "automated chain that ends with a rich Slack alert in approximately 5-7 seconds."
    )
    doc.add_paragraph()

    steps = [
        ("Trigger", "A CloudWatch Alarm fires (e.g., CPU > 80% for 2 consecutive 1-minute periods) OR an EC2 instance state changes (stopped/terminated/stopping) detected by EventBridge."),
        ("SNS Fan-out", "The alarm/event is published to the SNS topic aiops-alerts-dev. SNS simultaneously delivers to: (1) the Lambda function and (2) an optional email subscriber."),
        ("Lambda Invocation", "SNS synchronously invokes aiops-incident-processor-dev with the event payload wrapped in an SNS Records structure."),
        ("Event Parser", "The event_parser module unwraps the SNS message, detects the event type (CloudWatch alarm vs. EC2 state change), and normalizes it into a standard incident dict with: incident_id (UUID), timestamp, alarm_name, instance_id, event_type, alarm_reason, log_group, region, environment."),
        ("Log Fetcher", "Queries the CloudWatch Logs API for the relevant log group, fetching up to 100 log events from the last 15 minutes. Gracefully returns a placeholder if the log group does not exist."),
        ("Log Sanitizer", "Applies 8 regex-based redaction patterns to remove sensitive data before sending to AI: IPv4 addresses, AWS access keys (AKIA*), 40-character secrets, email addresses, internal hostnames (ip-*.internal), user home paths (/home/*/Users/*), and credential patterns (password=, token=, key=)."),
        ("Log Trimmer", "Intelligently caps logs at 4,000 characters (~1,000 tokens). Priority lines containing keywords like error, critical, fatal, exception, traceback, OOM, segfault, timeout, connection refused, disk full are always preserved. Remaining budget goes to the most recent lines."),
        ("Error Classifier", "A lightweight pre-scan classifies the error type from keywords in the alarm reason and logs: HIGH_CPU, MEMORY_EXHAUSTION, DISK_FULL, NETWORK_ISSUE, INSTANCE_FAILURE, APPLICATION_ERROR, or UNKNOWN. This seeds the AI prompt."),
        ("AI Root Cause Analysis", "A structured prompt is built and sent to Amazon Bedrock (Claude 3.5 Sonnet). If Bedrock fails, it falls back to OpenRouter, trying up to 5 free-tier models in sequence. The AI returns a structured JSON response with: summary, root_cause, severity (CRITICAL/HIGH/MEDIUM/LOW), severity_reason, affected_components, immediate_actions, long_term_fix, pattern_detected, confidence, and estimated_impact."),
        ("DynamoDB Persistence", "The enriched incident (metadata + AI analysis) is saved to DynamoDB table aiops-incidents-dev. A 30-day TTL attribute is set for automatic expiration. DynamoDB failures are logged but do NOT block the Slack notification."),
        ("Slack Alert", "A Block Kit formatted message is POST'd to the Slack webhook (fetched from Secrets Manager). The message includes severity-coded header, full incident metadata, root cause, immediate actions, long-term fix, pattern warning (if detected), confidence level, and estimated impact. Total time from alarm to Slack: ~5-7 seconds."),
    ]

    for i, (title_text, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        r_num = p.add_run(f"Step {i}: ")
        r_num.bold = True
        r_num.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        r_num.font.size = Pt(11)
        r_title = p.add_run(f"{title_text}  —  ")
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_desc = p.add_run(desc)
        r_desc.font.size = Pt(10)

    doc.add_page_break()

    # ── 4. TECH STACK ─────────────────────────────────────────────
    add_heading(doc, "4. Technology Stack", 1, "1F3864")
    add_divider(doc)
    add_table(doc,
        ["Layer", "Technology", "Details"],
        [
            ["Language", "Python 3.12", "All Lambda code; uses only stdlib + boto3 (no heavy deps)"],
            ["Infrastructure as Code", "Terraform 1.7+", "Modular HCL; S3 remote backend; modules for networking, EC2, IAM, Lambda, alarms"],
            ["CI/CD", "GitHub Actions", "4-stage pipeline: Lint → Unit Test → Terraform Apply → Lambda Deploy + Smoke Test"],
            ["Cloud Provider", "AWS (ap-south-1)", "Mumbai region; all services within single region"],
            ["Compute (Processing)", "AWS Lambda", "512 MB, 300s timeout, Python 3.12, X-Ray tracing, SQS DLQ"],
            ["Compute (Monitored)", "EC2 t3.micro + ASG", "Auto Scaling Group (1-3 instances), Launch Template, ALB for multi-AZ"],
            ["AI — Primary", "Amazon Bedrock\nClaude 3.5 Sonnet", "Invoked via boto3; 3 retries with exponential backoff; 2048 max tokens"],
            ["AI — Fallback", "OpenRouter\n(free tier models)", "Gemma 3N → Nemotron 9B → GPT-OSS 20B → Qwen 3 → GLM 4.5; 30s per model timeout"],
            ["Incident Storage", "Amazon DynamoDB", "Pay-per-request billing; 30-day TTL; incident_id + timestamp as composite key"],
            ["Log Archive", "Amazon S3", "AES-256 encrypted; public access blocked; lifecycle: IA → Glacier → delete"],
            ["Messaging", "Amazon SNS", "Topic-based fan-out to Lambda + email; EventBridge integration via resource policy"],
            ["Failure Handling", "Amazon SQS (DLQ)", "14-day message retention; captures all failed Lambda invocations"],
            ["Notifications", "Slack (Block Kit)", "Webhook via urllib (no requests library); severity-coded colors; rich structured format"],
            ["Secrets", "AWS Secrets Manager", "Slack webhook URL encrypted at rest; in-memory cached in Lambda runtime"],
            ["Monitoring", "CloudWatch + X-Ray", "4 alarms; live dashboard; 30-day log retention; distributed tracing on all invocations"],
            ["Networking", "VPC + ALB", "10.0.0.0/16; public subnets (ALB), private subnets (EC2); NAT Gateway for outbound"],
            ["Linting", "flake8 + flake8-bugbear", "Max line length 120; runs on every push/PR"],
            ["Testing", "pytest", "10 unit tests across event parser, log processor, and notification formatter; no AWS needed"],
        ]
    )

    doc.add_page_break()

    # ── 5. LAMBDA MODULE BREAKDOWN ────────────────────────────────
    add_heading(doc, "5. Lambda Internal Architecture", 1, "1F3864")
    add_divider(doc)
    add_para(doc,
        "There is a single Lambda function (aiops-incident-processor-dev) with a modular "
        "internal structure. Each module has a single responsibility and can be tested independently."
    )
    doc.add_paragraph()
    add_table(doc,
        ["Module", "File", "Responsibility"],
        [
            ["Handler (Orchestrator)", "incident_processor/handler.py", "Entry point. Runs all 6 pipeline steps in sequence. Handles fallbacks. Saves to DynamoDB. Re-raises on failure so DLQ captures it."],
            ["Event Parser", "incident_processor/event_parser.py", "Unwraps SNS envelope. Detects CloudWatch alarm vs EC2 state change. Returns normalized incident dict with UUID incident_id, timestamps, instance ID, log group, alarm details."],
            ["Log Fetcher", "incident_processor/log_fetcher.py", "Calls CloudWatch Logs filter_log_events API. 15-minute lookback. Up to 100 events. Returns formatted string. Graceful fallback for missing log groups."],
            ["Log Processor", "log_processor/processor.py", "Orchestrates sanitize → trim → classify. Builds the complete AI-ready payload dict."],
            ["Log Sanitizer", "log_processor/log_sanitizer.py", "8 regex redaction patterns. Removes IPs, AWS keys, secrets, emails, internal hostnames, user paths, and credential key-value pairs."],
            ["Log Trimmer", "log_processor/log_trimmer.py", "Hard cap: 4000 chars / 50 lines. Priority keywords (error, OOM, traceback, etc.) always kept. Fills remaining budget with most recent lines."],
            ["AI Analyzer", "ai_analyzer/analyzer.py", "Builds prompt, routes to Bedrock, falls back to OpenRouter on any exception. Returns enriched payload with ai_analysis dict."],
            ["Bedrock Client", "ai_analyzer/bedrock_client.py", "Invokes Bedrock via boto3. Supports Claude, Nova, and Titan request formats. 3 retries with backoff. Parses JSON from response (handles markdown fences)."],
            ["OpenRouter Client", "ai_analyzer/groq_client.py", "Calls OpenRouter API via urllib (no requests dep). Tries 5 free-tier models in sequence. 30s timeout per model. Same JSON parsing and validation."],
            ["RCA Prompt Builder", "ai/prompts/rca_prompt.py", "Constructs structured prompt: SRE persona + full incident context + sanitized logs + JSON-only output requirement with explicit schema."],
            ["Slack Notifier", "notification_handler/notifier.py", "Fetches webhook URL from Secrets Manager. Builds and POSTs Block Kit message via urllib. Logs failures but does not raise (non-blocking)."],
            ["Slack Formatter", "notification_handler/slack_formatter.py", "Builds full Block Kit JSON structure. Severity-based colors. Conditional pattern warning block. Handles missing fields gracefully."],
            ["Secrets Fetcher", "notification_handler/secrets.py", "Gets secret from Secrets Manager. Caches result in module-level variable for Lambda container reuse (avoids repeated API calls)."],
        ]
    )

    doc.add_page_break()

    # ── 6. INFRASTRUCTURE SUMMARY ─────────────────────────────────
    add_heading(doc, "6. Infrastructure Summary (Terraform)", 1, "1F3864")
    add_divider(doc)
    add_table(doc,
        ["Resource", "Name", "Key Configuration"],
        [
            ["Lambda Function", "aiops-incident-processor-dev", "Python 3.12 | 512 MB | 300s timeout | X-Ray active | SQS DLQ | SNS trigger"],
            ["SNS Topic", "aiops-alerts-dev", "Subscriptions: Lambda (primary) + Email (optional) | EventBridge publish policy"],
            ["EventBridge Rule", "aiops-ec2-state-change-dev", "Matches EC2 state: stopped / terminated / stopping → SNS"],
            ["DynamoDB Table", "aiops-incidents-dev", "PAY_PER_REQUEST | Hash: incident_id | Range: timestamp | 30-day TTL"],
            ["S3 Bucket (Logs)", "aiops-log-archive-dev", "AES-256 | Block all public access | Lifecycle: 30d→IA, 90d→Glacier, 365d→delete"],
            ["S3 Bucket (State)", "aiops-terraform-state-{account}", "Terraform remote backend | Versioning enabled"],
            ["SQS Dead Letter Queue", "aiops-lambda-dlq-dev", "14-day message retention | Captures failed Lambda invocations"],
            ["Secrets Manager", "aiops/slack/webhook", "Slack webhook URL | 7-day recovery | lifecycle ignore_changes"],
            ["VPC", "aiops-vpc-dev", "CIDR: 10.0.0.0/16 | DNS enabled | 2 AZs"],
            ["Public Subnets", "aiops-public-subnet-1/2-dev", "10.0.0.0/24 + 10.0.1.0/24 | ALB, NAT Gateway"],
            ["Private Subnets", "aiops-private-subnet-1/2-dev", "10.0.10.0/24 + 10.0.11.0/24 | EC2 instances"],
            ["Internet Gateway", "aiops-igw-dev", "Attached to VPC for public subnet outbound"],
            ["NAT Gateway", "aiops-nat-dev", "Private subnet outbound internet access"],
            ["Auto Scaling Group", "aiops-asg-dev", "Min: 1 | Max: 3 | Desired: 2 | t3.micro | CloudWatch detailed monitoring"],
            ["Launch Template", "aiops-lt-dev", "Amazon Linux 2023 | User data: install httpd + stress-ng for testing"],
            ["Application Load Balancer", "aiops-alb-dev", "Public | Multi-AZ | Port 80 health checks | Target: ASG"],
            ["CloudWatch Alarm (CPU High)", "aiops-high-cpu-dev", "CPU > 80% | 2 periods (1m avg) | Alarm → SNS"],
            ["CloudWatch Alarm (CPU Low)", "aiops-low-cpu-dev", "CPU < 20% | 3 periods (1m avg) | Informational"],
            ["CloudWatch Alarm (Status)", "aiops-status-check-failed-dev", "StatusCheckFailed >= 1 | 2 periods | Alarm → SNS"],
            ["CloudWatch Alarm (Network)", "aiops-network-in-high-dev", "NetworkIn > 50 MB | 2 periods (5m avg) | Alarm → SNS"],
            ["CloudWatch Dashboard", "AIOps-Sentinel-dev", "CPU/Memory gauges | Lambda KPIs | Latency percentiles | Alarm status tiles"],
            ["CloudWatch Log Group", "/aws/lambda/aiops-incident-processor-dev", "30-day retention | All Lambda stdout/stderr"],
            ["IAM Role (Lambda)", "aiops-lambda-role-dev", "Fine-grained: scoped to aiops-* resources only"],
            ["IAM Role (EC2)", "aiops-ec2-role-dev", "SSM Session Manager + CloudWatch agent permissions"],
        ]
    )

    doc.add_page_break()

    # ── 7. CI/CD PIPELINE ─────────────────────────────────────────
    add_heading(doc, "7. CI/CD Pipeline (GitHub Actions)", 1, "1F3864")
    add_divider(doc)
    add_para(doc,
        "Every push to main and every pull request triggers the CI/CD pipeline. "
        "Deploy jobs only run on main pushes and are skipped for commits starting with 'docs:' "
        "to avoid unnecessary infrastructure deploys for documentation changes."
    )
    doc.add_paragraph()

    add_table(doc,
        ["Job", "Trigger", "Steps", "Duration"],
        [
            ["1. Lint", "All pushes + PRs", "Checkout → Setup Python 3.12 → pip install flake8 → flake8 on lambda/, ai/, tests/ (max line 120)", "~30s"],
            ["2. Unit Tests", "All pushes + PRs\n(needs: lint)", "Checkout → Setup Python → pip install boto3 pytest → Run 3 test suites (10 tests total, no AWS needed)", "~60s"],
            ["3. Terraform Apply", "main push only\nnon-docs commits\n(needs: test)", "Checkout → AWS credentials → Setup Terraform 1.7.0 → init → validate → plan → apply", "~2 min"],
            ["4. Deploy Lambda", "main push only\nnon-docs commits\n(needs: deploy-infra)", "Checkout → AWS credentials → python scripts/deploy_lambda.py → verify config → smoke test (aws lambda invoke)", "~45s"],
        ]
    )
    doc.add_paragraph()
    add_heading(doc, "GitHub Secrets Required", 2, "2E74B5")
    add_table(doc,
        ["Secret", "Purpose"],
        [
            ["AWS_ACCESS_KEY_ID", "IAM user access key for CI/CD deployments"],
            ["AWS_SECRET_ACCESS_KEY", "IAM user secret key"],
            ["AWS_ACCOUNT_ID", "12-digit AWS account ID (652197206400)"],
            ["AMI_ID", "Amazon Linux 2023 AMI ID for ap-south-1"],
            ["ALERT_EMAIL", "Optional email for SNS subscription"],
            ["GROQ_API_KEY", "OpenRouter API key for AI fallback"],
        ]
    )
    doc.add_paragraph()
    add_heading(doc, "Lambda Deploy Script (scripts/deploy_lambda.py)", 2, "2E74B5")
    add_para(doc, "The deploy script performs 5 steps:")
    add_numbered(doc, "Creates a temp directory and copies all Lambda source files (lambda/, ai/prompts/)")
    add_numbered(doc, "Zips all .py files into deployment_package.zip")
    add_numbered(doc, "Uploads the ZIP to Lambda via boto3 update_function_code")
    add_numbered(doc, "Waits for the update to complete (polls function state)")
    add_numbered(doc, "Verifies deployment by checking LastUpdateStatus == Successful")

    doc.add_page_break()

    # ── 8. AI ANALYSIS DEEP DIVE ──────────────────────────────────
    add_heading(doc, "8. AI Analysis — Deep Dive", 1, "1F3864")
    add_divider(doc)

    add_heading(doc, "Primary: Amazon Bedrock (Claude 3.5 Sonnet)", 2, "2E74B5")
    add_para(doc, "Model: anthropic.claude-3-5-sonnet-20241022-v2:0 (configurable via Terraform variable)")
    add_para(doc, "Configuration:")
    add_bullet(doc, "Max tokens: 2,048 (optimized for cost vs. detail balance)")
    add_bullet(doc, "Temperature: 0.1 (near-deterministic — consistency over creativity)")
    add_bullet(doc, "Top-p: 0.9")
    add_bullet(doc, "Retries: 3 attempts with exponential backoff (2s, 4s, 6s) on throttling")
    add_bullet(doc, "Also supports: Amazon Nova Lite and Amazon Titan (auto-detected from model ID)")
    doc.add_paragraph()

    add_heading(doc, "Fallback: OpenRouter (Free Tier)", 2, "2E74B5")
    add_para(doc, "Triggered when Bedrock raises any exception. Tries these models in order:")
    add_numbered(doc, "google/gemma-3n-e4b-it:free")
    add_numbered(doc, "nvidia/nemotron-nano-9b-v2:free")
    add_numbered(doc, "openai/gpt-oss-20b:free")
    add_numbered(doc, "qwen/qwen3-4b:free")
    add_numbered(doc, "z-ai/glm-4.5-air:free")
    add_para(doc, "30-second timeout per model. If all 5 fail, the handler catches the exception and returns a HIGH-severity fallback response flagging manual review.")
    doc.add_paragraph()

    add_heading(doc, "RCA Prompt Structure", 2, "2E74B5")
    add_para(doc, "The prompt instructs the AI to act as an expert AWS SRE and provides:")
    add_bullet(doc, "Incident metadata: ID, timestamp, event type, alarm name/state/reason, instance ID, region, environment")
    add_bullet(doc, "Pre-classified error type from keyword scan")
    add_bullet(doc, "Sanitized and trimmed logs (last 15 minutes, max 4,000 chars)")
    add_bullet(doc, "Explicit JSON-only output requirement with field schema")
    add_bullet(doc, "Severity guidance: CRITICAL=service down, HIGH=major degradation, MEDIUM=performance impact, LOW=informational")
    doc.add_paragraph()

    add_heading(doc, "Sample AI Output", 2, "2E74B5")
    code_block = doc.add_paragraph()
    code_block.paragraph_format.left_indent = Inches(0.4)
    r = code_block.add_run(
        '{\n'
        '  "summary": "Java heap exhaustion caused OOM, halting all application threads",\n'
        '  "root_cause": "Unbounded cache growth in DataProcessor.java triggered heap exhaustion.\n'
        '                  GC overhead reached 98%, causing full GC every 2s. JVM halted threads.",\n'
        '  "severity": "HIGH",\n'
        '  "severity_reason": "Full service unavailability for ~5 min, required manual restart",\n'
        '  "affected_components": ["EC2", "Auto Scaling Group", "ALB"],\n'
        '  "immediate_actions": [\n'
        '    "Restart affected instance (ASG will auto-replace it)",\n'
        '    "Increase JVM heap to 4GB in launch template",\n'
        '    "Scale ASG to 3 instances during recovery"\n'
        '  ],\n'
        '  "long_term_fix": "Implement bounded LRU cache with 10k entry limit. Add heap alarm at 80%.",\n'
        '  "pattern_detected": true,\n'
        '  "pattern_description": "OOM crash every ~2h since v2.3.1 deployment",\n'
        '  "confidence": "HIGH",\n'
        '  "estimated_impact": "~5 min full outage, ~200 users affected"\n'
        '}'
    )
    r.font.name = "Courier New"
    r.font.size = Pt(9)

    doc.add_page_break()

    # ── 9. SECURITY & IAM ─────────────────────────────────────────
    add_heading(doc, "9. Security & IAM", 1, "1F3864")
    add_divider(doc)

    add_heading(doc, "IAM — Least Privilege", 2, "2E74B5")
    add_para(doc, "The Lambda execution role is scoped exclusively to aiops-* named resources. It cannot access any other DynamoDB table, S3 bucket, SNS topic, or secret in the account.")
    add_table(doc,
        ["Service", "Permissions Granted", "Scope"],
        [
            ["CloudWatch Logs", "CreateLogGroup, CreateLogStream, PutLogEvents, FilterLogEvents", "All log groups (needed to write own logs)"],
            ["Amazon Bedrock", "InvokeModel, InvokeModelWithResponseStream", "All models (AWS-managed — no resource-level scoping available)"],
            ["Amazon DynamoDB", "PutItem, GetItem, UpdateItem, Query, Scan", "aiops-incidents* tables only"],
            ["Amazon S3", "PutObject, GetObject, ListBucket", "aiops-log-archive* buckets only"],
            ["Secrets Manager", "GetSecretValue", "aiops/* secrets only"],
            ["Amazon SNS", "Publish", "aiops-* topics only"],
            ["Amazon SQS", "SendMessage, GetQueueAttributes", "aiops-* queues only"],
            ["Amazon EC2", "DescribeInstances, Start/Stop/RebootInstances", "Account-wide (required for remediation)"],
            ["CloudWatch", "GetMetricStatistics, DescribeAlarms", "Account-wide read-only"],
            ["AWS X-Ray", "PutTraceSegments, PutTelemetryRecords", "Account-wide (required for tracing)"],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "Secrets Management", 2, "2E74B5")
    add_bullet(doc, "Slack webhook URL stored in AWS Secrets Manager (aiops/slack/webhook)")
    add_bullet(doc, "Fetched at Lambda runtime via GetSecretValue API call")
    add_bullet(doc, "Cached in Lambda module-level variable — avoids repeated API calls on warm invocations")
    add_bullet(doc, "Terraform lifecycle: ignore_changes = [secret_string] — prevents accidental overwrite after manual update")
    add_bullet(doc, "OpenRouter API key stored as GitHub Secret and passed as Lambda env var via Terraform")
    doc.add_paragraph()

    add_heading(doc, "Log Sanitization (Pre-AI Redaction)", 2, "2E74B5")
    add_para(doc, "Before logs are sent to any external AI service, 8 redaction patterns are applied:")
    add_table(doc,
        ["Pattern", "Replacement", "Example"],
        [
            ["IPv4 addresses", "[IP_REDACTED]", "192.168.1.100 → [IP_REDACTED]"],
            ["AWS access keys (AKIA*)", "[AWS_KEY_REDACTED]", "AKIAIOSFODNN7EXAMPLE → [AWS_KEY_REDACTED]"],
            ["40-character secrets", "[AWS_SECRET_REDACTED]", "wJalrXUtnFEMI/K7MDENG/... → [AWS_SECRET_REDACTED]"],
            ["Email addresses", "[EMAIL_REDACTED]", "user@company.com → [EMAIL_REDACTED]"],
            ["Internal EC2 hostnames", "[INTERNAL_HOST_REDACTED]", "ip-10-0-1-5.ap-south-1.compute.internal → [INTERNAL_HOST_REDACTED]"],
            ["User home paths", "/home/[USER]/", "/home/ubuntu/app → /home/[USER]/app"],
            ["Credential key-value pairs", "[CREDENTIAL_REDACTED]", "password=secret123 → [CREDENTIAL_REDACTED]"],
            ["Token patterns", "[CREDENTIAL_REDACTED]", "token=abc123xyz → [CREDENTIAL_REDACTED]"],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "S3 & Network Security", 2, "2E74B5")
    add_bullet(doc, "S3 log archive: Block all public access enabled; AES-256 server-side encryption")
    add_bullet(doc, "EC2 instances in private subnets — not directly accessible from internet")
    add_bullet(doc, "ALB in public subnets — only entry point for HTTP traffic")
    add_bullet(doc, "Security groups restrict inbound to port 80 (ALB) and SSH/SSM only")
    add_bullet(doc, "SSM Session Manager used instead of SSH keys for EC2 access")

    doc.add_page_break()

    # ── 10. MONITORING ────────────────────────────────────────────
    add_heading(doc, "10. Monitoring & Observability", 1, "1F3864")
    add_divider(doc)

    add_heading(doc, "CloudWatch Alarms", 2, "2E74B5")
    add_table(doc,
        ["Alarm", "Metric", "Threshold", "Periods", "Action"],
        [
            ["aiops-high-cpu-dev", "CPUUtilization", "> 80%", "2 × 1 min avg", "Publish to SNS → trigger full pipeline"],
            ["aiops-low-cpu-dev", "CPUUtilization", "< 20%", "3 × 1 min avg", "Publish to SNS (informational, under-utilization)"],
            ["aiops-status-check-failed-dev", "StatusCheckFailed", ">= 1", "2 × 1 min max", "Publish to SNS → critical, treat missing as breaching"],
            ["aiops-network-in-high-dev", "NetworkIn", "> 50 MB/5 min", "2 × 5 min avg", "Publish to SNS → traffic spike / possible DDoS"],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, "CloudWatch Dashboard (AIOps-Sentinel-dev)", 2, "2E74B5")
    add_table(doc,
        ["Widget", "Type", "What It Shows"],
        [
            ["Title Banner", "Text", "Environment, region, account, 'Powered by AI-driven RCA'"],
            ["CPU Utilization", "Gauge", "Real-time CPU %, red line at alarm threshold (80%)"],
            ["Memory Utilization", "Gauge", "CWAgent mem_used_percent, red line at 85%"],
            ["Incidents Today", "Single Value + Sparkline", "Lambda invocation count (setPeriodToTimeRange)"],
            ["Pipeline Failures", "Single Value + Sparkline", "Lambda error count (setPeriodToTimeRange)"],
            ["Incidents Saved", "Single Value + Sparkline", "DynamoDB PutItem SampleCount (setPeriodToTimeRange)"],
            ["Avg Lambda Duration", "Single Value", "Average and P99 execution time in seconds"],
            ["CPU Trend", "Time Series", "Avg + Max CPU over time with alarm threshold line"],
            ["Lambda Invocations vs Errors", "Time Series", "Invocation count, error count, and calculated error rate %"],
            ["Lambda Latency P50/P90/P99", "Time Series", "Percentile latency with 30s warning line"],
            ["Recent Incidents Log Feed", "Log Insights Table", "Live feed: Pipeline complete / AI analysis / severity messages"],
            ["Active Alarms", "Alarm Status", "Visual alarm state tiles (green/red/yellow)"],
        ]
    )

    doc.add_page_break()

    # ── 11. ANTICIPATED Q&A ───────────────────────────────────────
    add_heading(doc, "11. Anticipated Follow-Up Questions", 1, "1F3864")
    add_divider(doc)

    qas = [
        (
            "Q: Why Lambda over ECS or a long-running service for processing?",
            "Lambda is ideal here because incidents are infrequent, bursty events (not continuous workloads). "
            "Lambda scales to zero cost when idle, scales instantly on demand, and has no server management overhead. "
            "The 300-second timeout is more than sufficient for the full pipeline. ECS would add ~30 seconds cold-start "
            "latency and persistent cost even at zero load."
        ),
        (
            "Q: What happens if Amazon Bedrock is unavailable or throttled?",
            "The Bedrock client retries up to 3 times with exponential backoff (2s, 4s, 6s) on throttling. "
            "If all retries fail, the AI Analyzer automatically falls back to OpenRouter and tries 5 free-tier "
            "models in sequence (Gemma, Nemotron, GPT-OSS, Qwen, GLM). If all AI options fail, the handler "
            "returns a predefined HIGH-severity response with a 'manual review required' flag, ensuring "
            "the Slack alert is always sent regardless of AI availability."
        ),
        (
            "Q: How do you prevent the same incident from being processed multiple times?",
            "Each incident gets a UUID (incident_id) generated at parse time. The DynamoDB record uses "
            "incident_id as the hash key, so re-processing the same event would overwrite the same record "
            "rather than create duplicates. SNS guarantees at-least-once delivery, so duplicate Slack "
            "notifications are possible in rare cases — a deduplication layer (DynamoDB conditional writes "
            "checking for existing incident_id) would be a future improvement."
        ),
        (
            "Q: What is the estimated cost to run this system?",
            "At low volume (dev environment, ~50 Lambda invocations/month): Lambda ~$0, DynamoDB ~$0 "
            "(free tier), S3 ~$0.01, CloudWatch ~$2-5/month for metrics/logs, EC2 t3.micro ~$15/month "
            "(if running 24/7), NAT Gateway ~$5-10/month. Total: approximately $25-35/month. "
            "Amazon Bedrock is the variable cost — Claude 3.5 Sonnet is ~$3/1M input tokens. "
            "With 2048 token prompts and ~50 incidents/month, Bedrock cost is < $1/month."
        ),
        (
            "Q: How do you handle false alarms or noisy alerts?",
            "CloudWatch alarms require the threshold to be breached for multiple consecutive evaluation "
            "periods (e.g., CPU > 80% for 2 consecutive 1-minute periods) before firing — reducing "
            "single-spike false positives. The AI analysis includes a 'pattern_detected' field that "
            "identifies recurring patterns vs. one-off spikes. Future improvement: add alarm state "
            "check in the event parser to skip 'OK' state transitions."
        ),
        (
            "Q: How is the Slack webhook URL secured?",
            "The webhook URL is stored in AWS Secrets Manager under 'aiops/slack/webhook' — never "
            "hardcoded or stored in environment variables directly. The Lambda IAM role has "
            "GetSecretValue permission scoped only to 'aiops/*' secrets. The secret is fetched "
            "at runtime and cached in a module-level variable for Lambda container reuse. "
            "Terraform's lifecycle ignore_changes ensures manual secret updates aren't overwritten "
            "by subsequent Terraform applies."
        ),
        (
            "Q: Can this scale to production or multiple environments?",
            "Yes — the Terraform code is fully modular and parameterized by environment variable. "
            "Creating a 'prod' environment requires a new terraform/environments/prod/ directory "
            "copying the dev configuration with different variable values. Resource names all include "
            "the environment suffix (aiops-*-prod). The CI/CD can be extended with environment-specific "
            "job conditions. The Lambda function scales automatically with SNS fanout."
        ),
        (
            "Q: What happens if the Lambda function itself times out (hits 300s)?",
            "If Lambda times out, AWS marks the invocation as a failure and sends the original SNS "
            "message to the SQS Dead Letter Queue (aiops-lambda-dlq-dev), where it's retained for "
            "14 days. The DLQ can be monitored via CloudWatch alarms on ApproximateNumberOfMessages, "
            "and messages can be manually replayed. The root cause of the timeout is most likely the "
            "OpenRouter fallback with 5 models × 30s timeout each (150s max) — the CI/CD smoke test "
            "now uses --cli-read-timeout 360 to handle this."
        ),
        (
            "Q: Why use OpenRouter free-tier models as fallback instead of a second Bedrock model?",
            "OpenRouter free-tier requires no additional AWS spend or provisioned capacity. Bedrock "
            "throttling errors typically indicate quota exhaustion — switching to a different Bedrock "
            "model would likely hit the same quota limits. OpenRouter provides genuine diversity: "
            "different infrastructure, different model families (Google, NVIDIA, OpenAI, Alibaba), "
            "and zero marginal cost. The tradeoff is that free-tier models are less capable than "
            "Claude 3.5 Sonnet, but for structured RCA output they are sufficient."
        ),
        (
            "Q: How are logs sanitized before being sent to external AI services?",
            "The log_sanitizer module applies 8 regex-based redaction patterns before any AI call. "
            "IPv4 addresses, AWS access keys (AKIA* pattern), 40-character base64 secrets, email "
            "addresses, EC2 internal hostnames, user home directory paths, and credential "
            "key-value pairs (password=, token=, key=, secret=) are all replaced with safe "
            "placeholder strings like [IP_REDACTED] and [CREDENTIAL_REDACTED]. This ensures "
            "no sensitive infrastructure data is transmitted to Bedrock or OpenRouter."
        ),
        (
            "Q: Who monitors the monitoring system? What if the Lambda fails silently?",
            "Multiple layers: (1) Lambda errors are captured in CloudWatch Metrics (Errors count) "
            "and visible on the dashboard. (2) Failed invocations are routed to the SQS DLQ. "
            "(3) CloudWatch Logs retain all Lambda output for 30 days. (4) X-Ray tracing captures "
            "every invocation for distributed trace analysis. (5) The CI/CD smoke test verifies "
            "Lambda invokes successfully after every deployment. A CloudWatch alarm on Lambda "
            "Errors or DLQ depth would be the natural next improvement."
        ),
        (
            "Q: How do you replay a failed incident event from the DLQ?",
            "Messages in the SQS DLQ (aiops-lambda-dlq-dev) retain the original SNS event payload "
            "for 14 days. To replay: (1) aws sqs receive-message to retrieve the message, "
            "(2) extract the Body field containing the original SNS notification, "
            "(3) aws lambda invoke with the payload to manually re-trigger the pipeline, "
            "(4) aws sqs delete-message to remove from DLQ after successful processing. "
            "An automated DLQ redrive policy targeting the original Lambda function can also be "
            "configured in SQS for hands-free replay."
        ),
    ]

    for q, a in qas:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        rq = p_q.add_run(q)
        rq.bold = True
        rq.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        rq.font.size = Pt(11)

        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.3)
        p_a.paragraph_format.space_after = Pt(6)
        ra = p_a.add_run(a)
        ra.font.size = Pt(10)

    doc.add_page_break()

    # ── APPENDIX: Key File Paths ───────────────────────────────────
    add_heading(doc, "Appendix: Key File Reference", 1, "1F3864")
    add_divider(doc)
    add_table(doc,
        ["File", "Purpose"],
        [
            ["lambda/incident_processor/handler.py", "Master pipeline orchestrator — entry point for all Lambda invocations"],
            ["lambda/incident_processor/event_parser.py", "CloudWatch alarm and EC2 state change event normalization"],
            ["lambda/incident_processor/log_fetcher.py", "CloudWatch Logs API — 15-minute lookback, 100-event cap"],
            ["lambda/log_processor/log_sanitizer.py", "8-pattern regex redaction before AI submission"],
            ["lambda/log_processor/log_trimmer.py", "Intelligent 4000-char cap preserving priority error lines"],
            ["lambda/ai_analyzer/bedrock_client.py", "Bedrock invocation with retry logic and multi-model format support"],
            ["lambda/ai_analyzer/groq_client.py", "OpenRouter fallback — 5 free-tier models in sequence"],
            ["lambda/notification_handler/slack_formatter.py", "Block Kit message builder with severity-based colors"],
            ["ai/prompts/rca_prompt.py", "Structured RCA prompt with full incident context and JSON schema"],
            ["terraform/environments/dev/main.tf", "Core AWS resources: SNS, DynamoDB, S3, Secrets Manager, EventBridge"],
            ["terraform/modules/lambda/main.tf", "Lambda function, DLQ, SNS subscription, CloudWatch logs"],
            ["terraform/modules/alarms/main.tf", "4 CloudWatch metric alarms with SNS actions"],
            ["terraform/environments/dev/dashboard.tf", "Full CloudWatch dashboard definition (12 widgets)"],
            [".github/workflows/cicd.yml", "4-stage CI/CD: Lint → Test → Terraform → Lambda deploy"],
            ["scripts/deploy_lambda.py", "Lambda packaging and deployment script (zip + upload + verify)"],
            ["tests/test_local_lambda.py", "Event parser unit tests (3 tests, no AWS)"],
            ["tests/test_log_processor.py", "Log sanitizer and trimmer tests (3 tests)"],
            ["tests/test_notifications.py", "Slack Block Kit formatter tests (4 tests)"],
        ]
    )

    # ── Save ──────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "AIOps_Sentinel_Review.docx")
    doc.save(output_path)
    print(f"\nDocument saved: {output_path}")
    print("Open it in Microsoft Word for the full formatted view.\n")


if __name__ == "__main__":
    build()
